'''

import io
import re
import os
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader
from bs4 import BeautifulSoup

# --- Helper: Clean whitespace ---
def clean_text(text: str) -> str:
    if not text:
        return ""
    cleaned_text = re.sub(r'\n\s*\n', '\n', text)
    return cleaned_text.strip()

# --- Helper: Strip HTML tags (for DOCX only) ---
def strip_html(html_string):
    if not html_string:
        return ""
    soup = BeautifulSoup(html_string, "html.parser")
    for br in soup.find_all("br"):
        br.replace_with("\n")
    return soup.get_text()

# --- DOCX GENERATION ---
def generate_docx_from_data(data):
    doc = Document()
    style = data.get('styleOptions', {})
    font_name = style.get('fontFamily', 'Calibri').split(',')[0]
    font_size = style.get('fontSize', 11)
    accent_color_hex = style.get('accentColor', '#34495e').lstrip('#')
    accent_color_rgb = RGBColor.from_string(accent_color_hex)

    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    heading_style = doc.styles.add_style('SectionHeading', 1)
    heading_style.font.name = font_name
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.font.color.rgb = accent_color_rgb
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)

    personal = data.get('personal', {})

    # Optional: Add letterhead/logo image
    if 'letterhead_image_b64' in style:
        logo_bytes = base64.b64decode(style['letterhead_image_b64'])
        with open("temp_logo.png", "wb") as f:
            f.write(logo_bytes)
        doc.add_picture("temp_logo.png", width=Inches(1.2))
        os.remove("temp_logo.png")

    # --- Header with name ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(personal.get('name', ''))
    runner.bold = True
    runner.font.name = font_name
    runner.font.size = Pt(24)
    runner.font.color.rgb = accent_color_rgb

    # --- Contact Info ---
    contact_items = [
        personal.get('email'),
        personal.get('phone'),
        personal.get('location')
    ]
    if personal.get('legalStatus') and personal.get('legalStatus') != 'Prefer not to say':
        contact_items.append(personal.get('legalStatus'))

    contact_info = " | ".join(filter(None, contact_items))
    p = doc.add_paragraph(contact_info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Optional: Profile picture ---
    if 'profile_image_b64' in personal:
        image_bytes = base64.b64decode(personal['profile_image_b64'])
        with open("temp_profile.png", "wb") as f:
            f.write(image_bytes)
        doc.add_picture("temp_profile.png", width=Inches(1.2))
        os.remove("temp_profile.png")

    doc.add_paragraph()

    # --- Summary ---
    if data.get('summary'):
        doc.add_paragraph('Summary', style='SectionHeading')
        doc.add_paragraph(clean_text(strip_html(data['summary'])), style=normal_style)

    # --- Experience ---
    if data.get('experience') and any(e.get('jobTitle') for e in data['experience']):
        doc.add_paragraph('Experience', style='SectionHeading')
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.add_run(exp.get('jobTitle', '')).bold = True
            p.add_run(f"\n{exp.get('company', '')} | {exp.get('dates', '')}\n").italic = True
            p.add_run(clean_text(strip_html(exp.get('description', ''))))
            p.style = normal_style
            p.paragraph_format.space_after = Pt(12)

    return doc

# --- PDF GENERATION ---
def generate_pdf_from_data(data):
    if data.get('summary'):
        data['summary'] = clean_text(data['summary'])

    if data.get('experience'):
        for exp in data['experience']:
            exp['jobTitle'] = clean_text(exp.get('jobTitle', ''))
            exp['company'] = clean_text(exp.get('company', ''))
            exp['dates'] = clean_text(exp.get('dates', ''))
            exp['description'] = clean_text(exp.get('description', ''))

    if data.get('education'):
        for edu in data['education']:
            edu['degree'] = clean_text(edu.get('degree', ''))
            edu['institution'] = clean_text(edu.get('institution', ''))
            edu['graduationYear'] = clean_text(edu.get('graduationYear', ''))
            edu['achievements'] = clean_text(edu.get('achievements', ''))

    if data.get('skills'):
        for skill in data['skills']:
            skill['category'] = clean_text(skill.get('category', ''))
            skill['skills_list'] = clean_text(skill.get('skills_list', ''))

    if data.get('certifications'):
        for cert in data['certifications']:
            cert['name'] = clean_text(cert.get('name', ''))
            cert['issuer'] = clean_text(cert.get('issuer', ''))
            cert['date'] = clean_text(cert.get('date', ''))

    # Load Jinja2 template from /templates
    env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), 'templates')))
    template = env.get_template('resume_template.html')
    rendered_html = template.render(**data)

    return HTML(string=rendered_html).write_pdf()
'''


import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader
from bs4 import BeautifulSoup

def clean_text(t): return re.sub(r'\n\s*\n','\n', t or '').strip()
def strip_html(s):
    soup = BeautifulSoup(s or '', 'html.parser')
    for br in soup.find_all('br'): br.replace_with('\n')
    return soup.get_text()

def generate_docx_from_data(data):
    doc = Document()
    # ── Logo ──
    if data.get('includeLogo'):
        logo = os.path.join(os.path.dirname(__file__), 'assets', 'PamTen_Logo.png')
        if os.path.exists(logo):
            doc.add_picture(logo, width=Inches(1.5))
            doc.add_paragraph()

    # ── Styles ──
    style = data.get('styleOptions', {})
    font = style.get('fontFamily','Calibri').split(',')[0]
    sz   = style.get('fontSize',11)
    accent = RGBColor.from_string(style.get('accentColor','#34495e').lstrip('#'))

    normal = doc.styles['Normal']
    normal.font.name = font
    normal.font.size = Pt(sz)

    heading = doc.styles.add_style('SectionHeading', 1)
    heading.font.name = font
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = accent
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after  = Pt(6)

    # ── Header ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data['personal'].get('name',''))
    r.bold = True; r.font.name = font; r.font.size = Pt(24); r.font.color.rgb = accent

    contact = " | ".join(filter(None, [
        data['personal'].get('email'),
        data['personal'].get('phone'),
        data['personal'].get('location')
    ]))
    p = doc.add_paragraph(contact)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Summary ──
    if data.get('summary'):
        doc.add_paragraph('Summary', style='SectionHeading')
        doc.add_paragraph(clean_text(strip_html(data['summary'])), style=normal)

    # ── Experience ──
    if data.get('experience'):
        doc.add_paragraph('Experience', style='SectionHeading')
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.add_run(exp.get('jobTitle','')).bold = True
            p.add_run(f"\n{exp.get('company','')} | {exp.get('dates','')}\n").italic = True
            p.add_run(clean_text(strip_html(exp.get('description',''))))
            p.style = normal
            p.paragraph_format.space_after = Pt(12)

    return doc

def generate_pdf_from_data(data):
    # ── Clean ──
    if data.get('summary'): data['summary']=clean_text(data['summary'])
    for section in ('experience','education'):
        for item in data.get(section,[]):
            for k in item: item[k]=clean_text(item[k])

    # ── Render HTML ──
    env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__),'assets')))
    tpl = env.get_template('resume_template.html')
    html = tpl.render(**data)

    # ── PDF ──
    return HTML(string=html,
                base_url=os.path.join(os.path.dirname(__file__),'assets')
               ).write_pdf()


